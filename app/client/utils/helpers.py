import os
import numpy as np
from docx import Document
from fpdf import FPDF
from pypdf import PdfReader
from llama_index.core.schema import Document as LLamaDocument
from utils.constant import FINAL_PROMPT, MAX_LENGTH, MULTI_QUERY_PROMPT, REPETITION_PENALTY, TEMPERATURE, TOP_K, TOP_P
from utils.style import text2html, chat_block


# class Custom_PDF(FPDF):
#     def __init__(self):
#         super().__init__()
#         self.add_font('Arial Unicode MS', '', "./font/arial-unicode-ms.ttf", uni=True)
#         self.set_font('Arial Unicode MS', '', 12)

class Custom_PDF(FPDF):
    def __init__(self):
        super().__init__()
        dir_path = os.path.dirname(os.path.realpath(__file__)) 
        font_path = os.path.join(dir_path, '..', 'font', 'arial-unicode-ms.ttf')
        self.add_font('Arial Unicode MS', '', font_path, uni=True)
        self.set_font('Arial Unicode MS', '', 12)


def export_chat_to_pdf(chat_history):
    pdf = Custom_PDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    for question, answer, _, _ in chat_history:
        pdf.cell(200, 10, txt=f"Question: {question}", ln=True)
        pdf.cell(200, 10, txt=f"Answer: {answer}", ln=True)

    pdf.output("chat_history.pdf")


def get_file_by_name(files, name):
    for file in files:
        if file.name == name:
            return file
    return None


def page_from_uploaded_file(files, file_name, p_id, st=None):
    page_content = None
    file = get_file_by_name(files, file_name)
    if file.type == "application/pdf":
        page_content = read_pdf(file)[p_id]
    elif file.type == "text/plain":
        page_content = read_plain_text(file)[p_id]
    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        page_content = read_word_document(file)[0]
    else:
        if st:
            st.error("Error in ===> \"page_from_uploaded_file()\"")
    return page_content


def page_from_local_file(file_path, p_id, st=None):
    page_content = None
    try:
        if file_path.split(".")[-1] == "pdf":
            texts = read_pdf(file_path)
            page_content = texts[p_id]
        elif file_path.split(".")[-1] == "docx":
            texts = read_word_document(file_path)
            page_content = texts[p_id]
        elif file_path.split(".")[-1] == "txt":
            page_content = read_plain_text(file_path)[0]
        else:
            if st:
                st.error("Error in ===> \"page_from_local_file()\"")

    except Exception as e:
        if st:
            st.error("Error in ===> \"page_from_local_file()\"")

    return page_content


def read_pdf(file):
    reader = PdfReader(file)
    texts = [page.extract_text() for page in reader.pages]
    return texts


def read_plain_text(file):
    texts = [file.read().decode("utf-8")]
    return texts

def read_word_document(file):
    doc = Document(file)
    texts = [para.text for para in doc.paragraphs]
    return texts

def documents_from_files(files, st=None):
    documents = []
    for file in files:
        if file.type == "application/pdf":
            texts = read_pdf(file)
        elif file.type == "text/plain":
            texts = read_plain_text(file)
        elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            texts = read_word_document(file)
        else:
            if st:
                st.error("Format de fichier non pris en charge.")

        if texts:
            documents += [LLamaDocument(text=text, metadata={"page_label": idx + 1,
                                                             "file_name": file.name,
                                                             "file_path": ""})
                          for idx, text in enumerate(texts)]
    return documents if documents else None


def init_session_state(st):
    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = []

    if 'question_index' not in st.session_state:
        st.session_state.question_index = 0

    if 'uploaded_files' not in st.session_state:
        st.session_state.uploaded_files = []

    if 'chroma_collection' not in st.session_state:
        st.session_state.chroma_collection = None

    if 'llm_client' not in st.session_state:
        st.session_state.llm_client = None

    if 'cross_encoder' not in st.session_state:
        st.session_state.cross_encoder = None

    if 'emb_fun' not in st.session_state:
        st.session_state.emb_fun = None

    if 'token_splitter' not in st.session_state:
        st.session_state.token_splitter = None

    if 'user_name' not in st.session_state:
        st.session_state.user_name = None

    if "is_new_upload" not in st.session_state:
        st.session_state.is_new_upload = False

    if "current_chroma_collection" not in st.session_state:
        st.session_state.current_chroma_collection = None

    if "chroma_client" not in st.session_state:
        st.session_state.chroma_client = None

    if "current_chroma_client" not in st.session_state:
        st.session_state.current_chroma_client = None
        

def generate_related_queries(query, llm_client):
    """
    Generate related queries using the new LLM client. to increase the diversity of the retrieved documents.
    """
    prompt = MULTI_QUERY_PROMPT.format(query=query)
    
    # Query the new LLM client
    response = llm_client.query({
            "inputs": prompt,
            "parameters": {
                "temperature": TEMPERATURE,
                "max_length": MAX_LENGTH,
                "top_p": TOP_P,
                "top_k": TOP_K, 
                "repetition_penalty": REPETITION_PENALTY, 
            }
        })
    
    # Extract the generated text from the response
    generated_text = response[0]['generated_text']
    
    # Split the generated text to get the part after the instruction
    generated_text = generated_text.split("[/INST]")[-1]
    
    # Split the text into individual queries
    related_queries = generated_text.split('\n')  
    
    # Remove any empty strings from the list of queries
    related_queries = [rq for rq in related_queries if rq]

    return related_queries


def retrieve_augmented_documents(query, aug_query, chroma_collection, n_results=10):
    queries = [query] + aug_query

    # retrieve docs for each query
    results = chroma_collection.query(query_texts=queries,
                                      n_results=n_results,
                                      include=['documents', 'embeddings', 'metadatas'])

    # Check if results are empty and return an empty set if so
    if not results['ids']:
        return set()

    # we are using a set to avoid duplicated docs (we assume that queries could retrieve the same docs sometimes)
    un_ret_ids = set()
    for ids in results['ids']:
        for id in ids:
            un_ret_ids.add(id)

    return un_ret_ids


def re_ranking_retrieved_documents(query, cross_encoder, ret_ids, chroma_collection, n_filtered=3):
    # retrieve all docs according to our augmented query
    results = chroma_collection.get(list(ret_ids))

    # constructed a pair of (query, document) to compute an advanced comparison between (query, documents)
    pairs = [[query, doc] for doc in results["documents"]]
    scores = cross_encoder(pairs)

    # sort pairs according to the scores and get indexes of 'n_filtered' first documents
    assert n_filtered >= 1 ; "make sure : n_filtered >= 1 !"
    idx = np.argsort(scores)[::-1][:n_filtered]

    # select our documents
    documents = "\n\n---------------\n\n".join([pairs[i][1] for i in idx])

    # get sources where we have retrieved these chunks
    sources = [results["metadatas"][i] for i in idx]

    return documents, sources


def rag_system(st, query, cross_encoder, llm_client, chroma_collection, n_results=4, n_filtered=2):
    # augment our query: by generating more related queries using the llm
    aug_query = generate_related_queries(query, llm_client)

    # retrieve all documents for each generated query
    ret_ids = retrieve_augmented_documents(query, aug_query, chroma_collection, n_results)

    # retrieve chunks from docs according to the query (similarity principal == cosine, ...)
    information, sources = re_ranking_retrieved_documents(query, cross_encoder, ret_ids, chroma_collection, n_filtered)
        
    prompt = FINAL_PROMPT.format(query=query, information=information)
    
    # Query the new LLM client
    answer = llm_client.query({
            "inputs": prompt,
            "parameters": {
                "temperature": 0.9,
                "max_length": 2048,
                "top_p": 0.9,
                "top_k": 30,
                "repetition_penalty": 1.1
            }
        })
    
    # Extract the generated text from the response
    answer = answer[0]['generated_text']
    
    # the answer is generated after [/INST] tag
    answer = str(answer.split("[/INST]")[-1])

    return answer, sources, information


def embedding_uploaded_files(uploaded_files, st):
    collection = st.session_state.current_chroma_collection

    documents = documents_from_files(uploaded_files, st)

    nodes = st.session_state.token_splitter.get_nodes_from_documents(documents)

    prev_id = 0
    ids = []
    docs = []
    metadata = []
    for idx, node in enumerate(nodes):

        if not node.text:
            continue

        ids.append(str(idx + prev_id))

        metadata.append({
            "id_": node.id_,
            **node.metadata,
            "start_char_idx": node.start_char_idx,
            "end_char_idx": node.end_char_idx,
            "metadata_seperator": node.metadata_seperator
        })

        docs.append(node.text)

    collection.add(ids=ids, documents=docs, metadatas=metadata)

    return collection


def re_create_collection(st, uploaded_files, cc_name):
    if uploaded_files:
        try:
            st.session_state.current_chroma_client.delete_collection(name=cc_name)
        except:
            pass
        st.session_state.current_chroma_collection = st.session_state.current_chroma_client.create_collection(
            name=cc_name,
            embedding_function=st.session_state.emb_fun)
        st.session_state.current_chroma_collection = embedding_uploaded_files(uploaded_files, st)


def display_chat_history(st, uploaded_files):
    with st.container():
        for q, a, sources, is_upload in st.session_state.chat_history:
            st.markdown(chat_block(st.session_state.user_name, q, is_question=True), unsafe_allow_html=True)
            st.markdown(chat_block("BenIsAlla EduBot", a, is_question=False), unsafe_allow_html=True)
            files = set([(src["file_path"], src["file_name"]) for src in sources])

            with st.container():
                st.markdown("<h6 style='color: #dcd914;'> Sources </h6>", unsafe_allow_html=True)
                for file_path, file_name in files:
                    for src in sources:
                        if src["file_name"] == file_name:
                            page_id = int(src["page_label"]) - 1  # pages start from [1]
                            page_content = (
                                page_from_uploaded_file(uploaded_files, file_name, page_id, st) if is_upload else
                                page_from_local_file(file_path, page_id, st)
                            )
                            start = page_content[0:int(src["start_char_idx"])]
                            body = page_content[int(src["start_char_idx"]): int(src["end_char_idx"])]
                            end = page_content[int(src["end_char_idx"]):]

                            page_content = text2html(page_content, "p")
                            page_content = f"""
                                {start} <span style='color: #dcd914;'> {body} </span> {end}
                            """

                            with st.expander(f"File: {file_name}  P:{page_id}"):
                                with st.container(height=300, border=False):
                                    st.markdown(page_content, unsafe_allow_html=True)


def query_component(st, is_upload):
    with st.form(key=f'query_form_{st.session_state.question_index}'):
        query = st.text_input(
            label="",
            label_visibility="collapsed",
            key=f'query_{st.session_state.question_index}',
            placeholder="Message BenIsAlla EduBot ...:")
        submit_button = st.form_submit_button(label='Send')

        collection = st.session_state.current_chroma_collection if is_upload else st.session_state.chroma_collection

        if submit_button and query:
            with st.spinner(text='Processing ...'):
                answer, sources, information = rag_system(st,
                                                          query=query,
                                                          cross_encoder=st.session_state.cross_encoder,
                                                          llm_client=st.session_state.llm_client,
                                                          chroma_collection=collection,
                                                          n_results=4,
                                                          n_filtered=2)
                st.session_state.chat_history.append((query, answer, sources, is_upload))

            st.session_state.question_index += 1
            st.rerun()


def upload_handler(st, uploaded_files, current_collection_name):
    if st.session_state.is_new_upload:
        with st.spinner(text='Processing Your Files ...'):
            re_create_collection(st, uploaded_files, current_collection_name)
            st.session_state.is_new_upload = False
